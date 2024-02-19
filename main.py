from docx import Document
import tkinter as tk
from tkinter import filedialog
from math import floor

# Create a root window and hide it
root = tk.Tk()
root.withdraw()

# Open the file dialog and get the selected directory
save_path = filedialog.askdirectory()

# Create a new Document
doc = Document()
order_no = input('Föy numarasını girin: ')
width = float(input('Eni girin: '))
max_height = float(input('Uzun kenarı girin: '))
min_height = float(input('Kısa kenarı girin: '))
slide_width = float(input('Slayt genişliğini girin: '))
case_height = float(input('Kasa yüksekliğini girin: '))

final_slide_with = slide_width - 1
final_max_height = max_height - case_height
slide_count = floor(width / final_slide_with)
slope = round((max_height - min_height) / slide_count, 2)

table = doc.add_table(rows=1, cols=4)

data = [
    ('Sipariş No:', order_no, 'Taban Genişliği:', width),
    ('Uzun Kenar:', max_height, 'Kısa Kenar:', min_height),
    ('Slayt Genişliği:', slide_width, 'Kalan Genişlik:', final_slide_with),
    ('Slayt Sayısı:', round(slide_count, 2), 'Slope Farkı:', slope)
]

for item in data:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = str(item[1])
    cells[2].text = item[2]
    cells[3].text = str(item[3])

doc.add_paragraph('')
doc.add_paragraph('Slayt Yükseklikleri:')

for i in range(int(slide_count)):
    doc.add_paragraph(f'{i + 1}. Slayt: {round(final_max_height, 1)}')
    final_max_height -= slope

# Save the document at the specified path
doc.save(f"{save_path}/{order_no}.docx")